import sys
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed.")
    sys.exit(1)

doc = Document()

# Title
title = doc.add_heading('Saheli - MVP Architecture & Technical Details', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    "Saheli is a mobile-first AI family assistant designed specifically for Indian households. "
    "It addresses the invisible 'cognitive load' burdening parents (particularly mothers) by centralizing "
    "and automating the tracking of medicines, family events (birthdays/anniversaries), and school updates."
)

# 2. Tech Stack Setup
doc.add_heading('2. Technology Stack', level=1)
stack = [
    ("Frontend (Client)", "Next.js 14, React.js, TypeScript, Tailwind CSS, custom Vanilla CSS tokens for Mobile-First native-app constraints."),
    ("Backend (API Server)", "FastAPI, Python 3.13, Pydantic, SQLAlchemy 2.0 (Async)."),
    ("Database", "PostgreSQL (provisioned via Supabase in the cloud) using the asyncpg driver."),
    ("AI / LLM Layer", "Groq API (Llama-3.3-70b-versatile) for lightning-fast inference, structured via OpenAI SDK compatibility. Google Gemini backend exists as a robust fallback."),
    ("State & API Communication", "Zustand for global state, Custom React Hooks wrapping Fetch calls to FastAPI endpoints.")
]
for name, desc in stack:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(name + ": ").bold = True
    p.add_run(desc)

# 3. Architecture & Data Flow
doc.add_heading('3. Architecture & Data Flow', level=1)
doc.add_paragraph("The application follows a decoupled Client-Server architecture:")
flows = [
    ("REST API Workflow", "Client Next.js components -> Fetch API -> FastAPI Router -> SQLAlchemy async Session -> PostgreSQL database."),
    ("AI Context-Injection Flow", "User sends a message -> Next.js calls /api/chat -> FastAPI constructs a system prompt injecting the user's current 'Family Context' (e.g., active medicine low-stocks, upcoming children's school events) -> Groq LLM generates response -> FastAPI returns JSON -> UI renders message bubble natively."),
    ("Database Migration Model", "Alembic handles database schema migrations, and SQLAlchemy's async engine manages connection pooling to the edge database.")
]
for f_name, f_desc in flows:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f_name + ": ").bold = True
    p.add_run(f_desc)

# 4. Core Modules (MVP)
doc.add_heading('4. Core MVP Modules', level=1)
modules = {
    "Dashboard": "Centralized hub showing 'Cognitive Time Saved' metrics, quick horizontal scrolling cards for pending tasks, and global Floating Action Button (FAB) for AI chat.",
    "Family Space": "Relationship tracker with bottom-sheet modals. Automatically categorizes upcoming birthdays with budget tags and allows instant 'Draft WhatsApp Message via AI' actions.",
    "Health Hub": "Medicine tracking dashboard. Tracks daily dosage, stock levels, and programmatically flags urgent refills with distinct visual alerts (e.g. '< 5 pills left').",
    "School Tracker": "Monitors children's academic schedules (PTMs, Sports Days, Fee Deadlines) and auto-tags tasks requiring parent financial or physical action.",
    "Saheli AI Chatbot": "A context-aware Conversational AI that acts as a 'digital elder sister (didi)'. It has access to the user's DB state and can converse naturally using Indian cultural context."
}
for mod, desc in modules.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(mod + ": ").bold = True
    p.add_run(desc)

# 5. UI/UX & Design System
doc.add_heading('5. UI/UX Design System', level=1)
doc.add_paragraph("The frontend utilizes a strict Mobile-First design language mimicking a native iOS/Android application. To achieve this on the web:")
ui = [
    "App Shell Constraint: Maximum viewport width is capped at 430px, centered horizontally on desktop screens to restrict layout distortion.",
    "Navigation: Desktop sidebars were entirely ripped out and replaced with a sticky mobile Bottom Tab Bar (Home, Family, Health, School, Chat) and a frosted-glass top App Bar.",
    "Color Palette: Warm, culturally resonant gradients. Primary themes use Plum (#5b2d8e), Rose (#c8456c), Teal (for Health metrics), and Amber (for Warnings).",
    "Aesthetics & Micro-interactions: Large touch-friendly targets, soft elevation shadows, swipeable horizontal scroll rows, and smooth pulse animations for the AI typing states."
]
for item in ui:
    doc.add_paragraph(item, style='List Bullet')

output_path = r'C:\Users\PC\Desktop\saheli\Saheli_MVP_Details.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
